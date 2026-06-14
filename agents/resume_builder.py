"""Render a tailored resume (structured sections) to a .docx file."""

import os

from docx import Document
from docx.shared import Pt

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "tailored_resumes")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def build_docx(tailored: dict, user_id: int, job_id: int) -> str:
    """Write a tailored resume to data/tailored_resumes/<user_id>_<job_id>.docx and return the path."""
    doc = Document()

    name = tailored.get("contact_name") or "Resume"
    doc.add_heading(name, level=0)

    summary = tailored.get("summary")
    if summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

    skills = tailored.get("skills") or []
    if skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(skills))

    experience = tailored.get("experience") or []
    if experience:
        doc.add_heading("Experience", level=1)
        for job in experience:
            title = job.get("title", "")
            company = job.get("company", "")
            dates = job.get("dates", "")
            header = doc.add_paragraph()
            run = header.add_run(f"{title} — {company}".strip(" —"))
            run.bold = True
            if dates:
                date_run = header.add_run(f"  ({dates})")
                date_run.italic = True
                date_run.font.size = Pt(10)
            for bullet in job.get("bullets") or []:
                doc.add_paragraph(bullet, style="List Bullet")

    education = tailored.get("education") or []
    if education:
        doc.add_heading("Education", level=1)
        for entry in education:
            doc.add_paragraph(entry, style="List Bullet")

    output_path = os.path.join(OUTPUT_DIR, f"{user_id}_{job_id}.docx")
    doc.save(output_path)
    return output_path
